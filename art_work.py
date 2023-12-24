import docx
import openpyxl
import pyspark
import os
import findspark
from docx.shared import Inches

findspark.init()
import pandas as pd
from pyspark.sql import SQLContext, SparkSession
from pyspark.sql.functions import *
import pandas as pd
import PySimpleGUI as sg
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from reportlab.pdfgen import canvas
import os
from pathlib import Path
from docxtpl import DocxTemplate
from docx2pdf import convert
from num2words import num2words
from openpyxl import load_workbook
from reportlab.lib.pagesizes import letter,landscape
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
import barcode
from barcode.writer import ImageWriter
from barcode import Code39
from barcode.writer import ImageWriter

#
# def read_file(input_file):
#     df = pd.read_excel(input_file)
#     mm=df.head()
#     # print(pdfFile+"/"+"outputs.pdf")
#     #
#     # # Display the contents of the DataFrame
#     # print(df)
#     #
#     # spark = SparkSession.builder.appName("NewApp").getOrCreate()
#     # print(spark.sparkContext.appName)
#     # # # Print all values in the DataFrame
#     # df=spark.read.csv(input_file,header=True,inferSchema=True)
#     # print(df)
#     # df.show()
#     # cc=df.select('CD')
#     # cc.show()
#     # df.createOrReplaceTempView("emp")
#     # pp=spark.sql("select name from emp")
#     # pp.show()
#     #
#     # for index, row in df.iterrows():
#     #     print(row)
#     # v=input_file
#     # df=DataFrame(v)
#     # df.show()
#     # try:
#
#     # with open(input_file, 'r', encoding='utf-8', errors='replace') as file:
#     #     content = file.read()
#     #     print(content)
#
#     # except Exception as e:
#     #     sg.popup_error(f"Error: {str(e)}")
#     #
#     # Create a blank white image
#     # width, height = 1200, 1200  # Adjust dimensions as needed
#     # background_color = (255, 255, 255)
#     # image = Image.new("RGB", (width, height), background_color)
#     # draw = ImageDraw.Draw(image)
#     #
#     # # Load a font
#     # font_size = 40
#     # font = ImageFont.truetype("arial.ttf", font_size)
#     #
#     # # Draw data onto the image
#     # text = "\n".join(cc)
#     # text_color = (0, 0, 0)  # Black
#     # text_position = (50, 50)
#     # draw.multiline_text(text_position, text, font=font, fill=text_color)
#     #
#     # # Save the image
#     # image.save("cd_label.png")
#     # print("Image created successfully.")
#
#
#     workbook=load_workbook(input_file)
#     print("book",workbook)
#     # workbook["hhgh"]
#     worksheet= workbook.active
#     # print(worksheet.cell("Printer-Name","dfdsr"))
#     max_row=worksheet.max_row
#     max_column=worksheet.max_column
#     print(max_column,"rows",max_row)
#     # pdfFile=pdfFile+"/"+"outputs.pdf"
#     # c=canvas.Canvas(pdfFile,pagesize=landscape(letter))
#     top_margin=3*inch
#     left_margin=0.25*inch
#     bottom_margin=0.25*inch
#     right_margin=0.25*inch
#     cell_width=(13.5*inch - left_margin -right_margin)/max_column
#     cell_height=(10*inch - top_margin -bottom_margin)/max_row
#     documentPath = r"C:\Users\ISLUSER\Desktop\amrith\old\Validator\proof.docx"
#     doc = DocxTemplate(documentPath)
#     a={}
#
#     for row in range(1,max_row+1):
#         for column in range(1,max_column+1):
#             cell=worksheet.cell(row=row,column=column)
#             print("columns",cell.value)
#
#
#     #         doc.render(cell.value)
#     # print(a)
#     # doc.save()
#
#
#
#     #         print(cell.value)
#     #
#     #         doc.render()
#     #         # pa = Path.cwd() / "qtn_.docx"
#     # doc.save(pa)
#     #         x=left_margin+(column-1)*cell_width
#     #         y=11*inch - (top_margin+row * cell_height)
#     #         print(int(x),int(y))
#     #
#     #         c.drawString(x, y,text)
#     #
#     # #
#     # c.save()
#
# def main():
#     sg.theme("DarkBlue")
#     layout = [
#         [sg.Text("Welcome:")],
#         [sg.Text("Enter the Input file")], [sg.Input(key="inFilePath", enable_events=True), sg.FileBrowse()],
#         # [sg.Text("Enter the Output file")], [sg.Input(key="OutFilePath", enable_events=True), sg.FolderBrowse()],
#         [sg.ProgressBar(100, orientation='h', size=(20, 20), key='progressbar', bar_color=('green', 'black'),
#                         visible=False)],
#         [sg.Button("OK"), sg.Button("Exit")]
#     ]
#
#     window = sg.Window("ARTWORK", layout, finalize=True)
#
#     while True:
#         event, values = window.read()
#         if event in (sg.WIN_CLOSED, "Exit"):
#             break
#         if event == "OK" and values["inFilePath"].strip():
#             successFlag = True
#             input_file = values["inFilePath"]
#             # output_file = values["OutFilePath"]
#             window['progressbar'].Update(visible=True)
#             read_file(input_file)
#             window['progressbar'].Update(visible=False)
#             window.disappear()
#             sg.popup("Success!!!!" if successFlag else "Failed!!!", title="Status", icon=r"R.ico")
#             window.reappear()
#
#     window.close()
#
#
# if __name__ == "__main__":
#     main()
#
#
# import openpyxl
# from docx import Document
#
# def read_excel_data(file_path, sheet_name):
#     workbook = openpyxl.load_workbook(file_path)
#     sheet = workbook[sheet_name]
#     data = {}
#     for row in sheet.iter_rows(min_row=1, values_only=True):
#         key = row[0]  # Assuming the first column contains keys
#         values = row[1:]  # Remaining columns as values
#         data[key] = values
#     return data
#
# def replace_placeholders_in_docx(doc, data):
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             text = run.text
#             for key, values in data.items():
#                 placeholder = f'{{{{{key}}}}}'
#                 if placeholder in text:
#                     # Replace the placeholder with the first value in the list
#                     text = text.replace(placeholder, str(values[0]))
#             run.text = text
#
# if __name__ == '__main__':
#     # Set your Excel file path and sheet name
#     EXCEL_FILE_PATH = r'C:\Users\ISLUSER\Desktop\amrith\old\Validator\CD_L.xlsx'
#     SHEET_NAME = 'Sheet1'
#
#     # Set your Word document file path
#     DOCX_FILE_PATH = r'C:\Users\ISLUSER\Desktop\amrith\old\Validator\proof.docx'
#
#     # Read data from Excel file
#     excel_data = read_excel_data(EXCEL_FILE_PATH, SHEET_NAME)
#
#     print(excel_data)
#     # Open the Word document
#     doc = Document(DOCX_FILE_PATH)
#
#     # Replace placeholders in the Word document
#     replace_placeholders_in_docx(doc, excel_data)
#
#     # Save the modified Word document
#     OUTPUT_DOCX_FILE_PATH = 'C:/Users/ISLUSER/Desktop/amrith/old/Validator/proof_modified.docx'
#     doc.save(OUTPUT_DOCX_FILE_PATH)
#
#     print(f"Word document with replaced values saved at: {OUTPUT_DOCX_FILE_PATH}")

from docx import Document

# def replace_placeholder(doc, placeholder, value):
#     # print("placeholder",placeholder,"values",value)
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             print(run.text)
#             if placeholder in run.text:
#                 print("placeholderss",placeholder)
#                 run.text = run.text.replace(placeholder, str(value))
                # print(run.text)


def insert_image(doc, variable_name, image_path):
    for paragraph in doc.paragraphs:
        if variable_name in paragraph.text:
            run = paragraph.runs[0]
            run.clear()
            run.add_picture(image_path, width=Inches(1.65), height=Inches(0.58))
            # Remove the variable_name text after inserting the image
            # print(paragraph.text)
            # run.add_picture(image_path, width=Inches(1.65),height=Inches(0.58))  # Adjust the width as needed


def read_file(input_file):
    df = pd.read_excel(input_file)
    c={}
    headers = df.columns.tolist()
    # length=len(headers)
    # for i in range(length):
    #     # c[headers[i]]=df[headers[i]].to_string(index=False)
    #     bb=df[headers[i]].tolist()
    #     lol_string = ''.join(map(str, bb))
    #     print(bb)
    #     c[headers[i]]=lol_string

    headers = df.columns.tolist()

    for header in headers:
        c[header] = ''.join(map(str, df[header].tolist()))

    print("dictionary",c)
    number=c["SW_Part_Number"]
    # mm=[i for i in number]
    length=len(number)
    segment_length = length // 12


    barcode_format = barcode.get_barcode_class('code39')

    my_barcode = barcode_format(number, writer=ImageWriter())
    # my_barcode.__str__().replace(number," ")
    pa = Path.cwd() / "{}_sw".format(number)
    print("name",pa.name)
    image_path=str(pa)+".png"
    print(image_path)
    image_path_saved=my_barcode.save(pa)
    # # If you want to set the text on the generated image separately
    #
    # # Save the modified image
    # image=Image.open(image_path_saved)
    # # Get a drawing context on the image
    # draw = ImageDraw.Draw(image)
    #
    # # Set the font and position to draw the text
    # font = None  # Replace with the font you want to use, e.g., ImageFont.load_default()
    # text_position = (10, 10)  # Adjust the position as needed
    #
    # # Draw the text on the image
    # draw.text(text_position, number, font=font)
    #
    # # Save the modified image
    # image.save(image_path_saved)
    # # c.update({"code":image_path_saved})
    print("after barcode",c)
    variable_name = "{{code}}"

    #
    # # #
    documentPath=r"C:\Users\ISLUSER\Desktop\amrith\old\Validator\proof.docx"
    # doc2 = DocxTemplate(documentPath)


    # Create an instance of a word document
    doc = Document(documentPath)


    insert_image(doc, variable_name, image_path)

    doc.save("artsv.docx")
    pa = Path.cwd() / "artsv.docx"
    doc2 = DocxTemplate(pa)

    doc2.render(c)
    #

    doc2.save(pa)


    # Create a PDF file
    pdf = canvas.Canvas("artwork.pdf")

    for paragraph in doc2.paragraphs:
        # Write each paragraph to the PDF
        pdf.drawString(100, 800, paragraph.text)
        pdf.showPage()

    # Save the PDF file
    pdf.save()

    convert(pa,"artwork.pdf")
    # #
    #
    # def insert_image(doc, variable_name, image_path,c):
    #     print(c)
    #     for paragraph in doc.paragraphs:
    #         if variable_name in paragraph.text:
    #             run = paragraph.runs[0]
    #             run.clear()
    #             run.add_picture(image_path, width=Inches(1.0))
    #             return c
    # pp=insert_image(doc,variable_name,image_path,c)
    # doc.add_picture(image_path)



    # # Add the picture to the paragraph
    # run = paragraph.add_run()
    # print("pic",c["code"])
    # run.add_picture(c["code"], width=Inches(2.0))  # Adjust width as needed


    # hed=df.iterrows()
    # print("iterrrr",hed)
    # print(headers)
    # a=headers[0]
    # print(a)
    # column_values = df[a].tolist()
    # print("column values",column_values)
    # # print(pdfFile+"/"+"outputs.pdf")
    #
    # # Display the contents of the DataFrame
    # print(df)
    #
    # spark = SparkSession.builder.appName("NewApp").getOrCreate()
    # print(spark.sparkContext.appName)
    # # # Print all values in the DataFrame
    # df=spark.read.csv(input_file,header=True,inferSchema=True)
    # print(df)
    # df.show()
    # cc=df.select('CD')
    # cc.show()
    # df.createOrReplaceTempView("emp")
    # pp=spark.sql("select name from emp")
    # pp.show()
    #
    # for index, row in df.iterrows():
    #     print(row)
    # v=input_file
    # df=DataFrame(v)
    # df.show()
    # try:

    # with open(input_file, 'r', encoding='utf-8', errors='replace') as file:
    #     content = file.read()
    #     print(content)

    # except Exception as e:
    #     sg.popup_error(f"Error: {str(e)}")
    #
    # Create a blank white image
    # width, height = 1200, 1200  # Adjust dimensions as needed
    # background_color = (255, 255, 255)
    # image = Image.new("RGB", (width, height), background_color)
    # draw = ImageDraw.Draw(image)
    #
    # # Load a font
    # font_size = 40
    # font = ImageFont.truetype("arial.ttf", font_size)
    #
    # # Draw data onto the image
    # text = "\n".join(cc)
    # text_color = (0, 0, 0)  # Black
    # text_position = (50, 50)
    # draw.multiline_text(text_position, text, font=font, fill=text_color)
    #
    # # Save the image
    # image.save("cd_label.png")
    # print("Image created successfully.")


    # workbook=load_workbook(input_file)
    # print("book",workbook)
    # worksheet= workbook.active
    # max_row=worksheet.max_row
    # max_column=worksheet.max_column
    # print(max_column,"rows",max_row)
    # # pdfFile=pdfFile+"/"+"outputs.pdf"
    # # c=canvas.Canvas(pdfFile,pagesize=landscape(letter))
    # top_margin=3*inch
    # left_margin=0.25*inch
    # bottom_margin=0.25*inch
    # right_margin=0.25*inch
    # cell_width=(13.5*inch - left_margin -right_margin)/max_column
    # cell_height=(10*inch - top_margin -bottom_margin)/max_row

    # for row in range(1,max_row+1):
    #     for column in range(1,max_column+1):
    #         cell=worksheet.cell(row=row,column=column)
    #         text=str(cell.value)
    #         print(cell.value)
    #         x=left_margin+(column-1)*cell_width
    #         y=11*inch - (top_margin+row * cell_height)
    #         print(int(x),int(y))
    #
    #         c.drawString(x, y,text)
    #
    # #
    # c.save()

def main():
    sg.theme("BluePurple")
    layout = [
        [sg.Text("Welcome...!")],
        [sg.Text("Enter the Input file")], [sg.Input(key="inFilePath", enable_events=True), sg.FileBrowse()],
        # [sg.Text("Enter the Output file")], [sg.Input(key="OutFilePath", enable_events=True), sg.FolderBrowse()],
        [sg.ProgressBar(100, orientation='h', size=(20, 20), key='progressbar', bar_color=('green', 'blue'),
                        visible=False)],
        [sg.Button("OK"), sg.Button("Exit")]
    ]

    window = sg.Window("ARTWORK", layout, finalize=True)
    window.set_icon(r"art.ico")

    while True:
        event, values = window.read()
        if event in (sg.WIN_CLOSED, "Exit"):
            break
        if event == "OK" and values["inFilePath"].strip():
            successFlag = True
            input_file = values["inFilePath"]
            # output_file = values["OutFilePath"]
            window['progressbar'].Update(visible=True)
            read_file(input_file)
            window['progressbar'].Update(visible=False)
            window.disappear()
            sg.popup("Done!!!!" if successFlag else "Failed!!!", title="Status", icon=r"dpattern.ico")
            window.reappear()

    window.close()


if __name__ == "__main__":
    main()
