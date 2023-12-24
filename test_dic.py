# # # import pandas as pd
# # # import matplotlib.pyplot as plt
# # #
# # # # Assuming df is your DataFrame
# # # df = pd.read_excel("CD_L.xlsx")
# # #
# # # def render_dataframe_to_pdf(dataframe, filename='outputsv.pdf'):
# # #     fig, ax = plt.subplots(figsize=(8, 6))  # Adjust the figure size as needed
# # #
# # #     ax.axis('tight')
# # #     ax.axis('off')
# # #
# # #     ax.table(cellText=dataframe.values, colLabels=dataframe.columns, cellLoc='center', loc='center')
# # #
# # #     plt.savefig(filename)
# # #     plt.close()
# # #
# # # render_dataframe_to_pdf(df)
# #
# #
# #
# # import pandas as pd
# # from reportlab.pdfgen import canvas
# #
# # # Read data from Excel
# # excel_file = "CD_L.xlsx"
# # df = pd.read_excel(excel_file)
# # print("data",df.head())
# #
# # # Function to render DataFrame to PDF with handling text overlap
# # def render_dataframe_to_pdf(dataframe, filename='outputo.pdf'):
# #     c = canvas.Canvas(filename, pagesize=(6872, 7792))  # Set pagesize for landscape mode
# #     width, height = c._pagesize
# #     margin = 50
# #
# #     # Calculate the available width for text
# #     available_width = width - 2 * margin
# #
# #     # Set font and font size
# #     c.setFont("Helvetica", 12)
# #
# #     # Calculate the column width based on the number of columns
# #     column_width = available_width / len(dataframe.columns)
# #
# #     # Set initial position for text
# #     x, y = margin, height - margin
# #
# #     # Iterate through each row and column to add text to the PDF
# #     for index, row in dataframe.iterrows():
# #         for col, value in row.items():
# #             # Check if the text exceeds the available width
# #             if x + column_width > width - margin:
# #                 # Move to the next row if the text exceeds the width
# #                 x = margin
# #                 y -= 10  # Adjust the spacing between rows
# #
# #             # Check for overlap with the bottom margin
# #             if y < margin:
# #                 # Move to the next page
# #                 c.showPage()
# #                 y = height - margin
# #
# #             c.drawString(x, y, f"{col}: {value}")
# #             x += column_width
# #
# #         # Move to the next row
# #         x = margin
# #         y -= 15  # Adjust the spacing between rows
# #
# #     c.save()
# #
# # # Call the function with your DataFrame and desired output filename
# # render_dataframe_to_pdf(df, filename='outputo.pdf')
#
#
#
# import pyspark
# import os
# import findspark
# findspark.init()
# import pandas as pd
# from pyspark.sql import SQLContext, SparkSession
# from pyspark.sql.functions import *
# import pandas as pd
# import PySimpleGUI as sg
# from pathlib import Path
# from PIL import Image, ImageDraw, ImageFont
#
# import os
#
# from openpyxl import load_workbook
# from reportlab.lib.pagesizes import letter,landscape
# from reportlab.lib.units import inch
# from reportlab.pdfgen import canvasdef read_file(input_file,pdfFile):
#     df = pd.read_excel(input_file)
#     mm=df.head()
#     print("bncbbcn",mm)
#     headers = df.columns.tolist()
#     print(headers)
#     print(pdfFile+"/"+"outputs.pdf")
#     #
#     # # Display the contents of the DataFrame
#     # print(df)
#     #
#     # spark = SparkSession.builder.appName("NewApp").getOrCreate()
#     # print(spark.sparkContext.appName)
#     # # # Print all values in the DataFrame
#     # df=spark.read.csv(input_file,header=True,inferSchema=True)
#     # print(df)
#     # df.show()
#     # cc=df.select('CD')
#     # cc.show()
#     # df.createOrReplaceTempView("emp")
#     # pp=spark.sql("select name from emp")
#     # pp.show()
#     #
#     # for index, row in df.iterrows():
#     #     print(row)
#     # v=input_file
#     # df=DataFrame(v)
#     # df.show()
#     # try:
#
#     # with open(input_file, 'r', encoding='utf-8', errors='replace') as file:
#     #     content = file.read()
#     #     print(content)
#
#     # except Exception as e:
#     #     sg.popup_error(f"Error: {str(e)}")
#     #
#     # Create a blank white image
#     # width, height = 1200, 1200  # Adjust dimensions as needed
#     # background_color = (255, 255, 255)
#     # image = Image.new("RGB", (width, height), background_color)
#     # draw = ImageDraw.Draw(image)
#     #
#     # # Load a font
#     # font_size = 40
#     # font = ImageFont.truetype("arial.ttf", font_size)
#     #
#     # # Draw data onto the image
#     # text = "\n".join(cc)
#     # text_color = (0, 0, 0)  # Black
#     # text_position = (50, 50)
#     # draw.multiline_text(text_position, text, font=font, fill=text_color)
#     #
#     # # Save the image
#     # image.save("cd_label.png")
#     # print("Image created successfully.")
#
#
#     workbook=load_workbook(input_file)
#     print("book",workbook)
#     worksheet= workbook.active
#     max_row=worksheet.max_row
#     max_column=worksheet.max_column
#     print(max_column,"rows",max_row)
#     pdfFile=pdfFile+"/"+"outputs.pdf"
#     c=canvas.Canvas(pdfFile,pagesize=landscape(letter))
#     top_margin=3*inch
#     left_margin=0.25*inch
#     bottom_margin=0.25*inch
#     right_margin=0.25*inch
#     cell_width=(13.5*inch - left_margin -right_margin)/max_column
#     cell_height=(10*inch - top_margin -bottom_margin)/max_row
#
#     for row in range(1,max_row+1):
#         for column in range(1,max_column+1):
#             cell=worksheet.cell(row=row,column=column)
#             text=str(cell.value)
#             print(cell.value)
#             x=left_margin+(column-1)*cell_width
#             y=11*inch - (top_margin+row * cell_height)
#             print(int(x),int(y))
#
#             c.drawString(x, y,text)
#
#     #
#     c.save()
#
# def main():
#     sg.theme("DarkBlue")
#     layout = [
#         [sg.Text("Welcome:")],
#         [sg.Text("Enter the Input file")], [sg.Input(key="inFilePath", enable_events=True), sg.FileBrowse()],
#         [sg.Text("Enter the Output file")], [sg.Input(key="OutFilePath", enable_events=True), sg.FolderBrowse()],
#         [sg.ProgressBar(100, orientation='h', size=(20, 20), key='progressbar', bar_color=('green', 'black'),
#                         visible=False)],
#         [sg.Button("OK"), sg.Button("Exit")]
#     ]
#
#     window = sg.Window("ARTWORK", layout, finalize=True)
#
#     while True:
#         event, values = window.read()
#         if event in (sg.WIN_CLOSED, "Exit"):
#             break
#         if event == "OK" and values["inFilePath"].strip():
#             successFlag = True
#             input_file = values["inFilePath"]
#             output_file = values["OutFilePath"]
#             window['progressbar'].Update(visible=True)
#             read_file(input_file,output_file)
#             window['progressbar'].Update(visible=False)
#             window.disappear()
#             sg.popup("Success!!!!" if successFlag else "Failed!!!", title="Status", icon=r"R.ico")
#             window.reappear()
#
#     window.close()
#
#
# if __name__ == "__main__":
#     main()
#
#
#



import barcode
from barcode.writer import ImageWriter
from openpyxl import Workbook
from openpyxl.drawing.image import Image
from pathlib import Path
from docxtpl import DocxTemplate
# import EAN13 from barcode module
from barcode import EAN13

# # import ImageWriter to generate an image file
# from barcode.writer import ImageWriter
#
# # Make sure to pass the number as string
# number = '499N4-60101'
#
# # Now, let's create an object of EAN13 class and
# # pass the number with the ImageWriter() as the
# # writer
# my_code = EAN13(number, writer=ImageWriter())
#
# # Our barcode is ready. Let's save it.
# my_code.save("new_code1")


import openpyxl
from openpyxl.drawing.image import Image
from barcode import Code128
from barcode.writer import ImageWriter

from barcode import Code128
from barcode.writer import ImageWriter
from barcode import Code39
# from barcode.writer import ImageWriter
# number='53P40-60101'
# barcode_format = barcode.get_barcode_class('code39')
#
#
# my_barcode =barcode_format(number, writer=ImageWriter())
# pa = Path.cwd() / "vv_.png"
# my_barcode.save(pa)


from docx import Document
from docx.shared import Inches
import barcode
from barcode.writer import ImageWriter

from docx import Document
from docx.shared import Inches

def insert_image(doc, variable_name, image_path):
    print(doc.paragraphs)
    for paragraph in doc.paragraphs:
        if variable_name in paragraph.text:
            run = paragraph.runs[0]
            run.clear()
            run.add_picture(image_path, width=Inches(1.0))  # Adjust the width as needed

# Load the template document
documentPath=r"C:\Users\ISLUSER\Desktop\amrith\old\Validator\proof.docx"
doc = Document(documentPath)

# Specify the variable name and image path
variable_name = "{{code}}"
image_path = r"C:\Users\ISLUSER\Desktop\amrith\old\Validator\vv_.png.png"

# Insert the image
insert_image(doc, variable_name, image_path)

# Save the modified document
doc.save("outputm.docx")


import pandas as pd
from docx import Document
from docx.shared import Inches
from pathlib import Path
import barcode
from barcode.writer import ImageWriter
#
# import pandas as pd
# from docx import Document
# from docx.shared import Inches
# from pathlib import Path
# import barcode
# from barcode.writer import ImageWriter
#
#
# def read_file(input_file):
#     df = pd.read_excel(input_file)
#     c = {}
#     headers = df.columns.tolist()
#     length = len(headers)
#
#     for i in range(length):
#         bb = df[headers[i]].tolist()
#         lol_string = ''.join(map(str, bb))
#         c[headers[i]] = lol_string
#
#     number = c["SW_Part_Number"]
#     barcode_format = barcode.get_barcode_class('code39')
#     my_barcode = barcode_format(number, writer=ImageWriter())
#
#     # Save the barcode image
#     image_path = Path.cwd() / "vvp_.png"
#     my_barcode.save(image_path)
#
#     # Update the dictionary with the barcode image path
#     c.update({"code": str(image_path)})
#
#     return c
#
#
# def insert_image(doc, variable_name, image_path):
#     for paragraph in doc.paragraphs:
#         if variable_name in paragraph.text:
#             run = paragraph.runs[0]
#             run.clear()
#             run.add_picture(image_path, width=Inches(1.0))
#
#
# def main():
#     input_file = "cd.xlsx"  # Replace with the path to your Excel file
#     document_path = r"C:\Users\ISLUSER\Desktop\amrith\old\Validator\proof.docx"
#
#     # Read data from Excel file and generate barcode
#     data_dict = read_file(input_file)
#
#     # Load the Word document
#     doc = Document(document_path)
#
#     # Insert the barcode image into the Word document
#     # insert_image(doc, data_dict)
#
#     # Save the modified Word document
#     # output_path = Path.cwd() / "qtn_.docx"
#     # doc.save(output_path)
#
#
# if __name__ == "__main__":
#     main()
#








import openpyxl
import pyspark
import os
import findspark
from docx.shared import Inches

findspark.init()
import pandas as pd
from pyspark.sql import SQLContext, SparkSession
from pyspark.sql.functions import *
import pandas as pd
import PySimpleGUI as sg
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont

import os
from pathlib import Path
from docxtpl import DocxTemplate
from docx2pdf import convert
from num2words import num2words
from openpyxl import load_workbook
from reportlab.lib.pagesizes import letter,landscape
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
import barcode
from barcode.writer import ImageWriter
from barcode import Code39
from barcode.writer import ImageWriter

#
# def read_file(input_file):
#     df = pd.read_excel(input_file)
#     mm=df.head()
#     # print(pdfFile+"/"+"outputs.pdf")
#     #
#     # # Display the contents of the DataFrame
#     # print(df)
#     #
#     # spark = SparkSession.builder.appName("NewApp").getOrCreate()
#     # print(spark.sparkContext.appName)
#     # # # Print all values in the DataFrame
#     # df=spark.read.csv(input_file,header=True,inferSchema=True)
#     # print(df)
#     # df.show()
#     # cc=df.select('CD')
#     # cc.show()
#     # df.createOrReplaceTempView("emp")
#     # pp=spark.sql("select name from emp")
#     # pp.show()
#     #
#     # for index, row in df.iterrows():
#     #     print(row)
#     # v=input_file
#     # df=DataFrame(v)
#     # df.show()
#     # try:
#
#     # with open(input_file, 'r', encoding='utf-8', errors='replace') as file:
#     #     content = file.read()
#     #     print(content)
#
#     # except Exception as e:
#     #     sg.popup_error(f"Error: {str(e)}")
#     #
#     # Create a blank white image
#     # width, height = 1200, 1200  # Adjust dimensions as needed
#     # background_color = (255, 255, 255)
#     # image = Image.new("RGB", (width, height), background_color)
#     # draw = ImageDraw.Draw(image)
#     #
#     # # Load a font
#     # font_size = 40
#     # font = ImageFont.truetype("arial.ttf", font_size)
#     #
#     # # Draw data onto the image
#     # text = "\n".join(cc)
#     # text_color = (0, 0, 0)  # Black
#     # text_position = (50, 50)
#     # draw.multiline_text(text_position, text, font=font, fill=text_color)
#     #
#     # # Save the image
#     # image.save("cd_label.png")
#     # print("Image created successfully.")
#
#
#     workbook=load_workbook(input_file)
#     print("book",workbook)
#     # workbook["hhgh"]
#     worksheet= workbook.active
#     # print(worksheet.cell("Printer-Name","dfdsr"))
#     max_row=worksheet.max_row
#     max_column=worksheet.max_column
#     print(max_column,"rows",max_row)
#     # pdfFile=pdfFile+"/"+"outputs.pdf"
#     # c=canvas.Canvas(pdfFile,pagesize=landscape(letter))
#     top_margin=3*inch
#     left_margin=0.25*inch
#     bottom_margin=0.25*inch
#     right_margin=0.25*inch
#     cell_width=(13.5*inch - left_margin -right_margin)/max_column
#     cell_height=(10*inch - top_margin -bottom_margin)/max_row
#     documentPath = r"C:\Users\ISLUSER\Desktop\amrith\old\Validator\proof.docx"
#     doc = DocxTemplate(documentPath)
#     a={}
#
#     for row in range(1,max_row+1):
#         for column in range(1,max_column+1):
#             cell=worksheet.cell(row=row,column=column)
#             print("columns",cell.value)
#
#
#     #         doc.render(cell.value)
#     # print(a)
#     # doc.save()
#
#
#
#     #         print(cell.value)
#     #
#     #         doc.render()
#     #         # pa = Path.cwd() / "qtn_.docx"
#     # doc.save(pa)
#     #         x=left_margin+(column-1)*cell_width
#     #         y=11*inch - (top_margin+row * cell_height)
#     #         print(int(x),int(y))
#     #
#     #         c.drawString(x, y,text)
#     #
#     # #
#     # c.save()
#
# def main():
#     sg.theme("DarkBlue")
#     layout = [
#         [sg.Text("Welcome:")],
#         [sg.Text("Enter the Input file")], [sg.Input(key="inFilePath", enable_events=True), sg.FileBrowse()],
#         # [sg.Text("Enter the Output file")], [sg.Input(key="OutFilePath", enable_events=True), sg.FolderBrowse()],
#         [sg.ProgressBar(100, orientation='h', size=(20, 20), key='progressbar', bar_color=('green', 'black'),
#                         visible=False)],
#         [sg.Button("OK"), sg.Button("Exit")]
#     ]
#
#     window = sg.Window("ARTWORK", layout, finalize=True)
#
#     while True:
#         event, values = window.read()
#         if event in (sg.WIN_CLOSED, "Exit"):
#             break
#         if event == "OK" and values["inFilePath"].strip():
#             successFlag = True
#             input_file = values["inFilePath"]
#             # output_file = values["OutFilePath"]
#             window['progressbar'].Update(visible=True)
#             read_file(input_file)
#             window['progressbar'].Update(visible=False)
#             window.disappear()
#             sg.popup("Success!!!!" if successFlag else "Failed!!!", title="Status", icon=r"R.ico")
#             window.reappear()
#
#     window.close()
#
#
# if __name__ == "__main__":
#     main()
#
#
# import openpyxl
# from docx import Document
#
# def read_excel_data(file_path, sheet_name):
#     workbook = openpyxl.load_workbook(file_path)
#     sheet = workbook[sheet_name]
#     data = {}
#     for row in sheet.iter_rows(min_row=1, values_only=True):
#         key = row[0]  # Assuming the first column contains keys
#         values = row[1:]  # Remaining columns as values
#         data[key] = values
#     return data
#
# def replace_placeholders_in_docx(doc, data):
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             text = run.text
#             for key, values in data.items():
#                 placeholder = f'{{{{{key}}}}}'
#                 if placeholder in text:
#                     # Replace the placeholder with the first value in the list
#                     text = text.replace(placeholder, str(values[0]))
#             run.text = text
#
# if __name__ == '__main__':
#     # Set your Excel file path and sheet name
#     EXCEL_FILE_PATH = r'C:\Users\ISLUSER\Desktop\amrith\old\Validator\CD_L.xlsx'
#     SHEET_NAME = 'Sheet1'
#
#     # Set your Word document file path
#     DOCX_FILE_PATH = r'C:\Users\ISLUSER\Desktop\amrith\old\Validator\proof.docx'
#
#     # Read data from Excel file
#     excel_data = read_excel_data(EXCEL_FILE_PATH, SHEET_NAME)
#
#     print(excel_data)
#     # Open the Word document
#     doc = Document(DOCX_FILE_PATH)
#
#     # Replace placeholders in the Word document
#     replace_placeholders_in_docx(doc, excel_data)
#
#     # Save the modified Word document
#     OUTPUT_DOCX_FILE_PATH = 'C:/Users/ISLUSER/Desktop/amrith/old/Validator/proof_modified.docx'
#     doc.save(OUTPUT_DOCX_FILE_PATH)
#
#     print(f"Word document with replaced values saved at: {OUTPUT_DOCX_FILE_PATH}")

from docx import Document

from docxtpl import DocxTemplate
from docx.shared import Inches
from pathlib import Path
import pandas as pd
import barcode
from barcode.writer import ImageWriter
#
#
# def insert_image(doc, variable_name, image_path):
#     for paragraph in doc.paragraphs:
#         if variable_name in paragraph.text:
#             run = paragraph.runs[0]
#             run.clear()
#             run.add_picture(image_path, width=Inches(1.0))
#
#
# def read_file(input_file):
#     df = pd.read_excel(input_file)
#     c = {}
#     headers = df.columns.tolist()
#
#     for header in headers:
#         c[header] = ''.join(map(str, df[header].tolist()))
#
#     number = c["SW_Part_Number"]
#     barcode_format = barcode.get_barcode_class('code39')
#     my_barcode = barcode_format(number, writer=ImageWriter())
#
#     pa = Path.cwd() / number
#     m = f"{pa}.png"
#     print(m)
#     image_path = my_barcode.save(pa)
#
#     variable_name = "{{code}}"
#     document_path = r"C:\Users\ISLUSER\Desktop\amrith\old\Validator\proof.docx"
#     doc = DocxTemplate(document_path)
#     doc2=Document(document_path)
#
#     insert_image(doc2, variable_name, image_path)
#     doc2.add_paragraph(c)
#     d=doc2.save("m.docx")
#     print(d)
#
#     doc.render(c)
#
#     output_path = Path.cwd() / "m.docx"
#     doc2.save(output_path)
#
#
#
#
#
#
#     # hed=df.iterrows()
#     # print("iterrrr",hed)
#     # print(headers)
#     # a=headers[0]
#     # print(a)
#     # column_values = df[a].tolist()
#     # print("column values",column_values)
#     # # print(pdfFile+"/"+"outputs.pdf")
#     #
#     # # Display the contents of the DataFrame
#     # print(df)
#     #
#     # spark = SparkSession.builder.appName("NewApp").getOrCreate()
#     # print(spark.sparkContext.appName)
#     # # # Print all values in the DataFrame
#     # df=spark.read.csv(input_file,header=True,inferSchema=True)
#     # print(df)
#     # df.show()
#     # cc=df.select('CD')
#     # cc.show()
#     # df.createOrReplaceTempView("emp")
#     # pp=spark.sql("select name from emp")
#     # pp.show()
#     #
#     # for index, row in df.iterrows():
#     #     print(row)
#     # v=input_file
#     # df=DataFrame(v)
#     # df.show()
#     # try:
#
#     # with open(input_file, 'r', encoding='utf-8', errors='replace') as file:
#     #     content = file.read()
#     #     print(content)
#
#     # except Exception as e:
#     #     sg.popup_error(f"Error: {str(e)}")
#     #
#     # Create a blank white image
#     # width, height = 1200, 1200  # Adjust dimensions as needed
#     # background_color = (255, 255, 255)
#     # image = Image.new("RGB", (width, height), background_color)
#     # draw = ImageDraw.Draw(image)
#     #
#     # # Load a font
#     # font_size = 40
#     # font = ImageFont.truetype("arial.ttf", font_size)
#     #
#     # # Draw data onto the image
#     # text = "\n".join(cc)
#     # text_color = (0, 0, 0)  # Black
#     # text_position = (50, 50)
#     # draw.multiline_text(text_position, text, font=font, fill=text_color)
#     #
#     # # Save the image
#     # image.save("cd_label.png")
#     # print("Image created successfully.")
#
#
#     # workbook=load_workbook(input_file)
#     # print("book",workbook)
#     # worksheet= workbook.active
#     # max_row=worksheet.max_row
#     # max_column=worksheet.max_column
#     # print(max_column,"rows",max_row)
#     # # pdfFile=pdfFile+"/"+"outputs.pdf"
#     # # c=canvas.Canvas(pdfFile,pagesize=landscape(letter))
#     # top_margin=3*inch
#     # left_margin=0.25*inch
#     # bottom_margin=0.25*inch
#     # right_margin=0.25*inch
#     # cell_width=(13.5*inch - left_margin -right_margin)/max_column
#     # cell_height=(10*inch - top_margin -bottom_margin)/max_row
#
#     # for row in range(1,max_row+1):
#     #     for column in range(1,max_column+1):
#     #         cell=worksheet.cell(row=row,column=column)
#     #         text=str(cell.value)
#     #         print(cell.value)
#     #         x=left_margin+(column-1)*cell_width
#     #         y=11*inch - (top_margin+row * cell_height)
#     #         print(int(x),int(y))
#     #
#     #         c.drawString(x, y,text)
#     #
#     # #
#     # c.save()
#
# def main():
#     sg.theme("DarkBlue")
#     layout = [
#         [sg.Text("Welcome:")],
#         [sg.Text("Enter the Input file")], [sg.Input(key="inFilePath", enable_events=True), sg.FileBrowse()],
#         # [sg.Text("Enter the Output file")], [sg.Input(key="OutFilePath", enable_events=True), sg.FolderBrowse()],
#         [sg.ProgressBar(100, orientation='h', size=(20, 20), key='progressbar', bar_color=('green', 'black'),
#                         visible=False)],
#         [sg.Button("OK"), sg.Button("Exit")]
#     ]
#
#     window = sg.Window("ARTWORK", layout, finalize=True)
#
#     while True:
#         event, values = window.read()
#         if event in (sg.WIN_CLOSED, "Exit"):
#             break
#         if event == "OK" and values["inFilePath"].strip():
#             successFlag = True
#             input_file = values["inFilePath"]
#             # output_file = values["OutFilePath"]
#             window['progressbar'].Update(visible=True)
#             read_file(input_file)
#             window['progressbar'].Update(visible=False)
#             window.disappear()
#             sg.popup("Success!!!!" if successFlag else "Failed!!!", title="Status", icon=r"R.ico")
#             window.reappear()
#
#     window.close()
#
#
# if __name__ == "__main__":
#     main()

import barcode
from barcode.writer import ImageWriter
from barcode import Code39
from barcode.writer import ImageWriter


def split_barcode(barcode, num_columns=13):
    # Calculate the length of each segment
    segment_length = len(barcode) // num_columns

    # Split the barcode into segments
    barcode_segments = [barcode[i:i + segment_length] for i in range(0, len(barcode), segment_length)]

    # If the length is not divisible by num_columns, the last segment may be shorter
    if len(barcode_segments) > num_columns:
        barcode_segments[-2] += barcode_segments[-1]
        barcode_segments.pop()

    return barcode_segments

# Example usage:
barcode = "1234567890123"  # Replace this with your actual barcode
columns = split_barcode(barcode, num_columns=13)

# Print the result
for i, column in enumerate(columns):
    print(f"Column {i + 1}: {column}")

